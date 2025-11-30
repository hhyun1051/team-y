"""
Create template DOCX files for delivery and product order documents
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Create delivery template (운송장)
doc = Document()

# Title
title = doc.add_heading('운송장', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph()

# Delivery info table
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

# Headers
headers = ['항목', '내용']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(12)

# Rows
rows_data = [
    ['수령인 이름', '{{NAME}}'],
    ['전화번호', '{{PHONE}}'],
    ['배송 주소', '{{ADDRESS}}']
]

for i, (label, placeholder) in enumerate(rows_data, start=1):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = placeholder

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('생성일: {{DATE}}')
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save('/root/office-worker/templates/delivery_template.docx')
print("[✅] Delivery template created: /root/office-worker/templates/delivery_template.docx")

# 2. Create product order template (거래명세서)
doc = Document()

# Title
title = doc.add_heading('거래명세서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph()

# Product info table
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'

# Headers
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(12)

# Rows
rows_data = [
    ['거래처', '{{CLIENT}}'],
    ['품목', '{{PRODUCT_NAME}}'],
    ['수량', '{{QUANTITY}}개'],
    ['단가', '{{UNIT_PRICE}}원'],
    ['합계', '{{TOTAL_PRICE}}원']
]

for i, (label, placeholder) in enumerate(rows_data, start=1):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = placeholder

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('생성일: {{DATE}}')
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save('/root/office-worker/templates/product_order_template.docx')
print("[✅] Product order template created: /root/office-worker/templates/product_order_template.docx")
