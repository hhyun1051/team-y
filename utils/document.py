"""
Document Processing Utilities

DOCX 템플릿 처리 및 PDF 변환 기능
"""

import os
import subprocess
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document


class DocumentProcessor:
    """DOCX 템플릿 처리 및 PDF 변환"""

    def __init__(self, templates_dir: str = "/root/office-worker/templates"):
        """
        DocumentProcessor 초기화

        Args:
            templates_dir: 템플릿 디렉토리 경로
        """
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def fill_template(
        self,
        template_name: str,
        data: Dict[str, Any],
        output_path: str
    ) -> str:
        """
        DOCX 템플릿에 데이터를 채워 새 문서 생성

        Args:
            template_name: 템플릿 파일명
            data: 채울 데이터 (dict)
            output_path: 출력 파일 경로

        Returns:
            생성된 DOCX 파일 경로
        """
        template_path = self.templates_dir / template_name

        if not template_path.exists():
            raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {template_path}")

        # 템플릿 로드
        doc = Document(template_path)

        # 텍스트 치환
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"  # {{key}} 형태의 플레이스홀더
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # 표(table) 내부 텍스트도 치환
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))

        # 파일 저장
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        print(f"[✅] DOCX 파일 생성 완료: {output_path}")
        return str(output_path)

    def convert_to_pdf(
        self,
        docx_path: str,
        pdf_path: Optional[str] = None
    ) -> str:
        """
        DOCX 파일을 PDF로 변환 (LibreOffice 사용)

        Args:
            docx_path: 변환할 DOCX 파일 경로
            pdf_path: 출력 PDF 경로 (None이면 자동 생성)

        Returns:
            생성된 PDF 파일 경로
        """
        docx_path = Path(docx_path)

        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX 파일을 찾을 수 없습니다: {docx_path}")

        # PDF 경로 설정
        if pdf_path is None:
            pdf_path = docx_path.with_suffix('.pdf')
        else:
            pdf_path = Path(pdf_path)

        # 출력 디렉토리 생성
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            # LibreOffice headless 모드로 PDF 변환
            # --headless: GUI 없이 실행
            # --convert-to pdf: PDF로 변환
            # --outdir: 출력 디렉토리
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent),
                str(docx_path)
            ]

            result = subprocess.run(
                cmd,
                check=True,
                capture_output=True,
                text=True,
                timeout=30
            )

            # LibreOffice는 입력 파일명과 같은 이름으로 PDF 생성
            # 원하는 파일명으로 변경
            generated_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"

            if generated_pdf != pdf_path and generated_pdf.exists():
                generated_pdf.rename(pdf_path)

            if not pdf_path.exists():
                raise RuntimeError(f"PDF 변환에 실패했습니다: {pdf_path}")

            print(f"[✅] PDF 변환 완료: {pdf_path}")
            return str(pdf_path)

        except subprocess.TimeoutExpired:
            raise RuntimeError("PDF 변환 시간 초과 (30초)")
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"PDF 변환 실패: {e.stderr}")
        except Exception as e:
            raise RuntimeError(f"PDF 변환 중 오류 발생: {str(e)}")

    def create_document_from_order(
        self,
        order_info: Dict[str, Any],
        template_name: str,
        output_dir: str = "/tmp/office-output"
    ) -> tuple[str, str]:
        """
        주문 정보로 DOCX 및 PDF 문서 생성

        Args:
            order_info: 주문 정보 (OrderInfo 또는 dict)
            template_name: 사용할 템플릿 이름
            output_dir: 출력 디렉토리

        Returns:
            tuple[str, str]: (DOCX 경로, PDF 경로)
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        # 출력 파일명 생성
        timestamp = os.popen("date +%Y%m%d_%H%M%S").read().strip()
        base_name = f"order_{timestamp}"

        docx_path = output_dir / f"{base_name}.docx"
        pdf_path = output_dir / f"{base_name}.pdf"

        # DOCX 생성
        docx_path = self.fill_template(
            template_name=template_name,
            data=order_info,
            output_path=str(docx_path)
        )

        # PDF 변환
        pdf_path = self.convert_to_pdf(
            docx_path=docx_path,
            pdf_path=str(pdf_path)
        )

        return docx_path, pdf_path
