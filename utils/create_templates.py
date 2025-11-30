"""
DOCX 템플릿 생성 스크립트

실행: python utils/create_templates.py
"""

from docx import Document
from pathlib import Path


def create_delivery_template():
    """배송 정보 템플릿 생성"""
    doc = Document()
    doc.add_heading('배송 정보', 0)

    # 표 생성
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # 표 내용
    cells = [
        ('항목', '내용'),
        ('수령인', '{{name}}'),
        ('전화번호', '{{phone}}'),
        ('주소', '{{address}}'),
    ]

    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # 저장
    template_dir = Path(__file__).parent.parent / 'templates'
    template_dir.mkdir(parents=True, exist_ok=True)
    output_path = template_dir / 'delivery_template.docx'
    doc.save(str(output_path))
    print(f"✅ 배송 정보 템플릿 생성: {output_path}")
    return output_path


def create_order_template():
    """제품 주문 템플릿 생성"""
    doc = Document()
    doc.add_heading('제품 주문서', 0)

    # 표 생성
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # 표 내용
    cells = [
        ('항목', '내용'),
        ('제품 종류', '{{product_type}}'),
        ('제원', '{{specifications}}'),
        ('수량', '{{quantity}}개'),
    ]

    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # 저장
    template_dir = Path(__file__).parent.parent / 'templates'
    template_dir.mkdir(parents=True, exist_ok=True)
    output_path = template_dir / 'order_template.docx'
    doc.save(str(output_path))
    print(f"✅ 제품 주문 템플릿 생성: {output_path}")
    return output_path


if __name__ == "__main__":
    print("[📝] DOCX 템플릿 생성 중...")
    create_delivery_template()
    create_order_template()
    print("[✅] 템플릿 생성 완료!")
