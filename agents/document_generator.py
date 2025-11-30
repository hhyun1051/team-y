"""
Document Generation Utilities

Handles DOCX template filling and PDF conversion using LibreOffice
"""
import os
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Dict, Any
from docx import Document


class DocumentGenerator:
    """문서 생성 및 PDF 변환 유틸리티"""

    TEMPLATE_DIR = Path("/root/office-worker/templates")
    OUTPUT_DIR = Path("/tmp")

    @staticmethod
    def fill_template(template_path: Path, replacements: Dict[str, str], output_path: Path) -> Path:
        """
        DOCX 템플릿에 데이터를 채워서 저장

        Args:
            template_path: 템플릿 DOCX 파일 경로
            replacements: 치환할 데이터 (key: placeholder, value: actual value)
            output_path: 출력 DOCX 파일 경로

        Returns:
            생성된 DOCX 파일 경로
        """
        doc = Document(template_path)

        # 단락의 텍스트 치환
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # 각 run에서 치환 (포맷 유지)
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        # 테이블의 텍스트 치환
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

        # 저장
        doc.save(output_path)
        print(f"[✅] DOCX generated: {output_path}")
        return output_path

    @staticmethod
    def convert_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
        """
        LibreOffice를 사용하여 DOCX를 PDF로 변환

        Args:
            docx_path: 입력 DOCX 파일 경로
            pdf_path: 출력 PDF 파일 경로 (디렉토리만 지정 가능)

        Returns:
            생성된 PDF 파일 경로
        """
        # LibreOffice headless 모드로 PDF 변환
        output_dir = pdf_path.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_dir),
            str(docx_path)
        ]

        print(f"[🔄] Converting DOCX to PDF: {docx_path.name}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr}")

        # LibreOffice는 같은 이름으로 PDF를 생성함
        expected_pdf = output_dir / f"{docx_path.stem}.pdf"

        if not expected_pdf.exists():
            raise FileNotFoundError(f"PDF not found: {expected_pdf}")

        # 원하는 이름으로 변경 (필요한 경우)
        if expected_pdf != pdf_path:
            expected_pdf.rename(pdf_path)

        print(f"[✅] PDF generated: {pdf_path}")
        return pdf_path

    @classmethod
    def generate_delivery_document(cls, name: str, phone: str, address: str) -> Dict[str, Path]:
        """
        배송 문서 생성 (DOCX + PDF)

        Args:
            name: 수령인 이름
            phone: 전화번호
            address: 배송 주소

        Returns:
            {"docx": Path, "pdf": Path}
        """
        template_path = cls.TEMPLATE_DIR / "delivery_template.docx"

        # 고유한 파일명 생성 (타임스탬프)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = cls.OUTPUT_DIR / f"delivery_{timestamp}.docx"
        pdf_path = cls.OUTPUT_DIR / f"delivery_{timestamp}.pdf"

        # 템플릿 치환 데이터
        replacements = {
            "{{NAME}}": name,
            "{{PHONE}}": phone,
            "{{ADDRESS}}": address,
            "{{DATE}}": datetime.now().strftime("%Y년 %m월 %d일"),
        }

        # DOCX 생성
        cls.fill_template(template_path, replacements, docx_path)

        # PDF 변환
        cls.convert_to_pdf(docx_path, pdf_path)

        return {"docx": docx_path, "pdf": pdf_path}

    @classmethod
    def generate_product_order_document(
        cls,
        client: str,
        product_name: str,
        quantity: int,
        unit_price: int
    ) -> Dict[str, Path]:
        """
        제품 주문 문서 생성 (DOCX + PDF)

        Args:
            client: 거래처
            product_name: 품목
            quantity: 수량
            unit_price: 단가

        Returns:
            {"docx": Path, "pdf": Path}
        """
        template_path = cls.TEMPLATE_DIR / "product_order_template.docx"

        # 고유한 파일명 생성 (타임스탬프)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = cls.OUTPUT_DIR / f"product_order_{timestamp}.docx"
        pdf_path = cls.OUTPUT_DIR / f"product_order_{timestamp}.pdf"

        # 합계 계산
        total_price = quantity * unit_price

        # 템플릿 치환 데이터
        replacements = {
            "{{CLIENT}}": client,
            "{{PRODUCT_NAME}}": product_name,
            "{{QUANTITY}}": str(quantity),
            "{{UNIT_PRICE}}": f"{unit_price:,}",
            "{{TOTAL_PRICE}}": f"{total_price:,}",
            "{{DATE}}": datetime.now().strftime("%Y년 %m월 %d일"),
        }

        # DOCX 생성
        cls.fill_template(template_path, replacements, docx_path)

        # PDF 변환
        cls.convert_to_pdf(docx_path, pdf_path)

        return {"docx": docx_path, "pdf": pdf_path}
