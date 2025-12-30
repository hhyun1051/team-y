"""
Document Generation Utilities

Handles DOCX template filling and PDF conversion using LibreOffice
"""
import os
import subprocess
import smtplib
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
from pdf2image import convert_from_path
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from dotenv import load_dotenv

# 환경 변수 로드
load_dotenv()


class DocumentGenerator:
    """문서 생성 및 PDF 변환 유틸리티"""

    # Docker 환경을 고려한 경로 설정
    TEMPLATE_DIR = Path(os.getenv("TEMPLATE_DIR", "/app/templates"))
    OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "/app/output"))

    # HP ePrint 설정 (환경 변수에서 로드)
    HP_PRINTER_EMAIL = os.getenv("HP_PRINTER_EMAIL")
    HP_SENDER_EMAIL = os.getenv("HP_SENDER_EMAIL")
    HP_SENDER_PASSWORD = os.getenv("HP_SENDER_PASSWORD")
    HP_SMTP_SERVER = os.getenv("HP_SMTP_SERVER", "smtp.gmail.com")
    HP_SMTP_PORT = int(os.getenv("HP_SMTP_PORT", "587"))

    @staticmethod
    def fill_template(template_path: Path, replacements: Dict[str, str], output_path: Path) -> Path:
        """
        DOCX 템플릿에 데이터를 채워서 저장

        Args:
            template_path: 템플릿 DOCX 파일 경로
            replacements: 치환할 데이터 (key: placeholder, value: actual value)
            output_path: 출력 DOCX 파일 경로

        Returns:
            생성된 DOCX 파일 경로
        """
        doc = Document(template_path)

        # 단락의 텍스트 치환
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # 각 run에서 치환 (포맷 유지)
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        # 테이블의 텍스트 치환 (스타일 보존)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 전체 단락 텍스트 확인
                        full_text = paragraph.text
                        needs_replacement = any(key in full_text for key in replacements.keys())

                        if needs_replacement:
                            # 모든 run의 텍스트를 교체 (스타일은 각 run에 유지됨)
                            for run in paragraph.runs:
                                original_text = run.text
                                new_text = original_text
                                for key, value in replacements.items():
                                    new_text = new_text.replace(key, value)
                                if new_text != original_text:
                                    run.text = new_text

        # 저장
        doc.save(output_path)
        print(f"[✅] DOCX generated: {output_path}")
        return output_path

    @staticmethod
    def convert_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
        """
        LibreOffice를 사용하여 DOCX를 PDF로 변환

        Args:
            docx_path: 입력 DOCX 파일 경로
            pdf_path: 출력 PDF 파일 경로 (디렉토리만 지정 가능)

        Returns:
            생성된 PDF 파일 경로
        """
        # LibreOffice headless 모드로 PDF 변환
        output_dir = pdf_path.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_dir),
            str(docx_path)
        ]

        print(f"[🔄] Converting DOCX to PDF: {docx_path.name}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr}")

        # LibreOffice는 같은 이름으로 PDF를 생성함
        expected_pdf = output_dir / f"{docx_path.stem}.pdf"

        if not expected_pdf.exists():
            raise FileNotFoundError(f"PDF not found: {expected_pdf}")

        # 원하는 이름으로 변경 (필요한 경우)
        if expected_pdf != pdf_path:
            expected_pdf.rename(pdf_path)

        print(f"[✅] PDF generated: {pdf_path}")
        return pdf_path

    @staticmethod
    def convert_to_images(pdf_path: Path, output_dir: Path = None, dpi: int = 150) -> List[Path]:
        """
        PDF를 이미지(PNG)로 변환

        Args:
            pdf_path: 입력 PDF 파일 경로
            output_dir: 출력 디렉토리 (None이면 PDF와 같은 디렉토리)
            dpi: 이미지 해상도 (기본값: 150, 높을수록 선명하지만 파일 크기 증가)

        Returns:
            생성된 이미지 파일 경로 리스트
        """
        if output_dir is None:
            output_dir = pdf_path.parent

        output_dir.mkdir(parents=True, exist_ok=True)

        print(f"[🔄] Converting PDF to images: {pdf_path.name}")

        try:
            # PDF를 이미지로 변환 (pdf2image 사용)
            images = convert_from_path(
                str(pdf_path),
                dpi=dpi,
                fmt='png',
                thread_count=2
            )

            image_paths = []
            base_name = pdf_path.stem

            for i, image in enumerate(images, start=1):
                # 이미지 파일명: <base_name>_page_<번호>.png
                image_path = output_dir / f"{base_name}_page_{i}.png"
                image.save(str(image_path), 'PNG')
                image_paths.append(image_path)
                print(f"[✅] Image generated: {image_path.name}")

            return image_paths

        except Exception as e:
            print(f"[❌] Image conversion failed: {str(e)}")
            return []

    @classmethod
    def print_pdf_to_hp(cls, pdf_path: Path, subject: str = "Print Document") -> bool:
        """
        HP ePrint를 사용하여 PDF를 프린터로 전송

        Args:
            pdf_path: 인쇄할 PDF 파일 경로
            subject: 이메일 제목 (기본값: "Print Document")

        Returns:
            성공 여부 (True/False)
        """
        # 프린터 설정 확인
        if not cls.HP_PRINTER_EMAIL or not cls.HP_SENDER_EMAIL or not cls.HP_SENDER_PASSWORD:
            print(f"[⚠️] HP ePrint not configured. Skipping print.")
            print(f"[ℹ️] Please set HP_PRINTER_EMAIL, HP_SENDER_EMAIL, HP_SENDER_PASSWORD in .env")
            return False

        # 파일 존재 확인
        if not pdf_path.exists():
            print(f"[❌] PDF file not found: {pdf_path}")
            return False

        try:
            print(f"[🖨️] Sending PDF to HP ePrint: {pdf_path.name}")

            # 이메일 메시지 생성
            msg = MIMEMultipart()
            msg['From'] = cls.HP_SENDER_EMAIL
            msg['To'] = cls.HP_PRINTER_EMAIL
            msg['Subject'] = subject

            # PDF 첨부
            with open(pdf_path, 'rb') as f:
                attachment = MIMEApplication(f.read(), _subtype='pdf')
                attachment.add_header(
                    'Content-Disposition',
                    'attachment',
                    filename=pdf_path.name
                )
                msg.attach(attachment)

            # SMTP로 전송
            with smtplib.SMTP(cls.HP_SMTP_SERVER, cls.HP_SMTP_PORT) as server:
                server.starttls()
                server.login(cls.HP_SENDER_EMAIL, cls.HP_SENDER_PASSWORD)
                server.send_message(msg)

            print(f"[✅] PDF sent to printer successfully: {pdf_path.name}")
            return True

        except smtplib.SMTPAuthenticationError as e:
            print(f"[❌] SMTP Authentication failed: {e}")
            print(f"[ℹ️] Please check HP_SENDER_EMAIL and HP_SENDER_PASSWORD")
            return False
        except smtplib.SMTPException as e:
            print(f"[❌] SMTP error: {e}")
            return False
        except Exception as e:
            print(f"[❌] Failed to send PDF to printer: {e}")
            return False

    @classmethod
    def print_pdfs_to_hp(cls, pdf_paths: List[Path], subject: str = "Print Documents") -> bool:
        """
        여러 PDF를 HP ePrint로 일괄 전송

        Args:
            pdf_paths: 인쇄할 PDF 파일 경로 리스트
            subject: 이메일 제목

        Returns:
            성공 여부 (True/False)
        """
        # 프린터 설정 확인
        if not cls.HP_PRINTER_EMAIL or not cls.HP_SENDER_EMAIL or not cls.HP_SENDER_PASSWORD:
            print(f"[⚠️] HP ePrint not configured. Skipping print.")
            return False

        try:
            print(f"[🖨️] Sending {len(pdf_paths)} PDF(s) to HP ePrint...")

            # 이메일 메시지 생성
            msg = MIMEMultipart()
            msg['From'] = cls.HP_SENDER_EMAIL
            msg['To'] = cls.HP_PRINTER_EMAIL
            msg['Subject'] = subject

            # PDF 첨부
            count = 0
            for pdf_path in pdf_paths:
                if pdf_path.exists():
                    with open(pdf_path, 'rb') as f:
                        attachment = MIMEApplication(f.read(), _subtype='pdf')
                        attachment.add_header(
                            'Content-Disposition',
                            'attachment',
                            filename=pdf_path.name
                        )
                        msg.attach(attachment)
                        count += 1
                        print(f"[📎] Attached: {pdf_path.name}")

            if count == 0:
                print(f"[⚠️] No valid PDF files to print")
                return False

            # SMTP로 전송
            with smtplib.SMTP(cls.HP_SMTP_SERVER, cls.HP_SMTP_PORT) as server:
                server.starttls()
                server.login(cls.HP_SENDER_EMAIL, cls.HP_SENDER_PASSWORD)
                server.send_message(msg)

            print(f"[✅] {count} PDF(s) sent to printer successfully")
            return True

        except Exception as e:
            print(f"[❌] Failed to send PDFs to printer: {e}")
            return False

    @classmethod
    def generate_delivery_document(
        cls,
        unloading_site: str,
        address: str,
        contact: str,
        payment_type: str,
        loading_site: str = "유진알루미늄",
        loading_address: str = None,
        loading_phone: str = None,
        freight_cost: int = None,
        notes: str = None,
        auto_print: bool = False
    ) -> Dict[str, Any]:
        """
        운송장 문서 생성 (DOCX + PDF)

        Args:
            unloading_site: 하차지 (회사 이름)
            address: 주소 (상세 주소)
            contact: 연락처
            payment_type: 운송비 지불 방법 ("착불" 또는 "선불")
            loading_site: 상차지 (기본값: "유진알루미늄")
            loading_address: 상차지 주소 (선택)
            loading_phone: 상차지 전화번호 (선택)
            freight_cost: 운송비 (착불일 경우에만, 원 단위)
            notes: 비고 (선택)
            auto_print: HP ePrint 자동 인쇄 여부 (기본값: False)

        Returns:
            {"docx": Path, "pdf": Path, "images": List[Path], "printed": bool}
        """
        template_path = cls.TEMPLATE_DIR / "deliver_template_new.docx"

        # 고유한 파일명 생성 (타임스탬프)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = cls.OUTPUT_DIR / f"delivery_{timestamp}.docx"
        pdf_path = cls.OUTPUT_DIR / f"delivery_{timestamp}.pdf"

        # 운송비 표시: 착불이고 금액이 있으면 금액 표시, 선불/착불(금액없음)이면 빈칸
        if payment_type == "착불" and freight_cost:
            freight_display = f"{freight_cost:,}원"
        else:
            freight_display = ""

        # 템플릿 치환 데이터
        replacements = {
            "{{UNLOADING_SITE}}": unloading_site,
            "{{ADDRESS}}": address,
            "{{CONTACT}}": contact,
            "{{LOADING_SITE}}": loading_site,
            "{{LOADING_ADDRESS}}": loading_address or "",
            "{{LOADING_PHONE}}": loading_phone or "",
            "{{PAYMENT_TYPE}}": payment_type,
            "{{FREIGHT_COST}}": freight_display,
            "{{NOTES}}": notes or "",
            "{{DATE}}": datetime.now().strftime("%Y년 %m월 %d일"),
        }

        # DOCX 생성
        cls.fill_template(template_path, replacements, docx_path)

        # PDF 변환
        cls.convert_to_pdf(docx_path, pdf_path)

        # 이미지 생성
        image_paths = cls.convert_to_images(pdf_path)

        # 자동 인쇄 (옵션)
        printed = False
        if auto_print:
            print(f"[🖨️] Auto-printing enabled for delivery document")
            printed = cls.print_pdf_to_hp(
                pdf_path,
                subject=f"운송장 - {unloading_site} ({timestamp})"
            )

        return {
            "docx": docx_path,
            "pdf": pdf_path,
            "images": image_paths,
            "printed": printed
        }

    @classmethod
    def generate_product_order_document(
        cls,
        client: str,
        product_name: str,
        quantity: int,
        unit_price: int,
        auto_print: bool = False
    ) -> Dict[str, Any]:
        """
        제품 주문 문서 생성 (DOCX + PDF)

        Args:
            client: 거래처
            product_name: 품목
            quantity: 수량
            unit_price: 단가
            auto_print: HP ePrint 자동 인쇄 여부 (기본값: False)

        Returns:
            {"docx": Path, "pdf": Path, "images": List[Path], "printed": bool}
        """
        template_path = cls.TEMPLATE_DIR / "product_order_template.docx"

        # 고유한 파일명 생성 (타임스탬프)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = cls.OUTPUT_DIR / f"product_order_{timestamp}.docx"
        pdf_path = cls.OUTPUT_DIR / f"product_order_{timestamp}.pdf"

        # 합계 계산
        total_price = quantity * unit_price

        # 템플릿 치환 데이터
        replacements = {
            "{{CLIENT}}": client,
            "{{PRODUCT_NAME}}": product_name,
            "{{QUANTITY}}": str(quantity),
            "{{UNIT_PRICE}}": f"{unit_price:,}",
            "{{TOTAL_PRICE}}": f"{total_price:,}",
            "{{DATE}}": datetime.now().strftime("%Y년 %m월 %d일"),
        }

        # DOCX 생성
        cls.fill_template(template_path, replacements, docx_path)

        # PDF 변환
        cls.convert_to_pdf(docx_path, pdf_path)

        # 이미지 생성
        image_paths = cls.convert_to_images(pdf_path)

        # 자동 인쇄 (옵션)
        printed = False
        if auto_print:
            print(f"[🖨️] Auto-printing enabled for product order document")
            printed = cls.print_pdf_to_hp(
                pdf_path,
                subject=f"거래명세서 - {client} ({timestamp})"
            )

        return {
            "docx": docx_path,
            "pdf": pdf_path,
            "images": image_paths,
            "printed": printed
        }